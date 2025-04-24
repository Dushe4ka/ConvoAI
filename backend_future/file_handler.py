# file_handler.py
from fastapi import UploadFile, HTTPException, Depends, File, Form, status
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
import os
import fitz # PyMuPDF
try:
    # python-docx для .docx файлов
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    DocxDocument = None
    PackageNotFoundError = None # Определяем, чтобы except работал
import uuid
import logging
from datetime import datetime
from typing import List, Dict, Optional, Any, Generator
from pathlib import Path
import chardet # Для определения кодировки .txt
import io # Для чтения .docx из памяти

# --- Импорты из проекта ---
from qdrant_db import ( # Используем функции-геттеры и классы моделей
    get_qdrant_client, get_embedding_model, PointStruct
)
from database import get_db, Document as DbDocument, Chat
from chat_AI import llm # Импортируем инициализированный LLM из chat_AI
# --- Импорт настроек и констант ---
from config import (
    GENERAL_CHAT_SESSION_ID, GENERAL_QDRANT_COLLECTION,
    CATEGORIES, MAX_FILE_SIZE, SUPPORTED_EXTENSIONS,
    CHUNK_SIZE, CHUNK_OVERLAP, UPLOAD_DIR
)

logger = logging.getLogger(__name__)

# Убедимся, что папка для временных файлов существует (хотя стараемся не использовать)
os.makedirs(UPLOAD_DIR, exist_ok=True)

class FileUploadError(Exception):
    """Кастомное исключение для ошибок обработки и загрузки файлов."""
    pass

# --- Функции Обработки Файлов ---

def get_file_encoding(content: bytes) -> str:
    """Определяет кодировку текстового файла по его содержимому."""
    detected = chardet.detect(content[:10000]) # Проверяем первые 10KB
    encoding = detected['encoding'] or 'utf-8' # По умолчанию utf-8
    confidence = detected['confidence']
    logger.debug(f"Определена кодировка: {encoding} с уверенностью {confidence}")
    return encoding

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> Generator[str, None, None]:
    """
    Разбивает текст на чанки заданного размера с перекрытием.
    Использует генератор для экономии памяти.
    Простая стратегия разбиения по предложениям/абзацам.
    """
    if not text:
        return

    # Разбиваем на параграфы или большие блоки
    blocks = text.split('\n\n')
    current_chunk = ""
    separator = ". " # Используем точку как разделитель предложений внутри чанка

    for block in blocks:
        block = block.strip()
        if not block:
            continue

        # Разбиваем блок на предложения (упрощенно)
        sentences = [s.strip() + "." for s in block.split('.') if s.strip()]

        for sentence in sentences:
            if len(current_chunk) + len(separator) + len(sentence) <= chunk_size:
                current_chunk += (separator if current_chunk else "") + sentence
            else:
                # Если текущий чанк не пуст, выдаем его
                if current_chunk:
                     # Добавляем перекрытие: начало текущего предложения
                     overlap_part = sentence[:overlap]
                     yield current_chunk # Отдаем предыдущий чанк
                     current_chunk = current_chunk[-overlap:] + separator + sentence # Начинаем новый с перекрытием
                else:
                     # Если предложение само по себе длиннее chunk_size, разбиваем его
                     for i in range(0, len(sentence), chunk_size - overlap):
                         yield sentence[i : i + chunk_size]
                     current_chunk = "" # Сбрасываем после разбивки большого предложения

    # Выдаем последний оставшийся чанк
    if current_chunk:
        yield current_chunk


async def process_file(file: UploadFile) -> Dict[str, Any]:
    """
    Обрабатывает загруженный UploadFile: проверяет, извлекает текст.
    Читает файл в память.
    """
    file_ext = Path(file.filename).suffix.lower()
    logger.info(f"Начало обработки файла: {file.filename}, тип: {file.content_type}, расширение: {file_ext}")

    # 1. Проверка расширения
    if file_ext not in SUPPORTED_EXTENSIONS:
        raise FileUploadError(f"Неподдерживаемый формат файла: {file_ext}. Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS.keys())}")
    if file_ext == '.docx' and DocxDocument is None:
         raise FileUploadError("Обработка .docx невозможна: библиотека python-docx не установлена.")

    # 2. Чтение файла в память и проверка размера
    try:
        content = await file.read()
        file_size = len(content)
        logger.debug(f"Файл '{file.filename}' прочитан в память, размер: {file_size} байт.")
    except Exception as e:
        logger.error(f"Ошибка чтения файла '{file.filename}' в память: {e}")
        raise FileUploadError(f"Ошибка чтения файла: {e}")
    finally:
        await file.close() # Закрываем файл после чтения

    if file_size == 0:
        raise FileUploadError("Файл пустой.")
    if file_size > MAX_FILE_SIZE:
        raise FileUploadError(f"Файл '{file.filename}' слишком большой ({file_size / (1024*1024):.2f} MB). Максимальный размер: {MAX_FILE_SIZE / (1024*1024):.0f} MB.")

    # 3. Извлечение текста в зависимости от типа файла
    text = ""
    try:
        if file_ext == '.pdf':
            with fitz.open(stream=content, filetype="pdf") as doc:
                text = "\n".join(page.get_text("text", sort=True) for page in doc).strip() # sort=True для лучшего порядка
                metadata = doc.metadata
                logger.debug(f"PDF Metadata: {metadata}")
        elif file_ext == '.docx':
             with io.BytesIO(content) as doc_stream:
                 doc = DocxDocument(doc_stream)
                 text = "\n".join(para.text for para in doc.paragraphs if para.text).strip()
        elif file_ext == '.txt':
            encoding = get_file_encoding(content)
            logger.info(f"Определена кодировка для '{file.filename}': {encoding}")
            try:
                text = content.decode(encoding).strip()
            except UnicodeDecodeError:
                 logger.warning(f"Не удалось декодировать '{file.filename}' как {encoding}, пробую utf-8 с игнорированием ошибок.")
                 text = content.decode('utf-8', errors='ignore').strip()

        if not text:
            # Возможно, файл содержит только изображения или защищен
            logger.warning(f"Не удалось извлечь текст из файла '{file.filename}'. Файл может быть пустым, содержать только изображения или быть защищенным.")
            # Не считаем это ошибкой, но чанков не будет
            # raise FileUploadError("Не удалось извлечь текст из файла (пустое или нетекстовое содержимое)")

        logger.info(f"Успешно извлечен текст из '{file.filename}', длина: {len(text)} символов.")
        return {
            "text": text,
            "original_filename": file.filename,
            "file_type": SUPPORTED_EXTENSIONS[file_ext],
            "file_size": file_size
        }

    except PackageNotFoundError: # Обработка ошибки, если .docx не является валидным zip-архивом
         logger.error(f"Файл '{file.filename}' не является валидным DOCX файлом.")
         raise FileUploadError("Файл поврежден или не является документом DOCX.")
    except Exception as e:
        logger.error(f"Ошибка извлечения текста из файла '{file.filename}': {e}", exc_info=True)
        raise FileUploadError(f"Ошибка обработки содержимого файла: {str(e)}")


def detect_category(text: str, filename: str) -> str:
    """Определяет категорию текста с помощью LLM."""
    if not llm:
        logger.error("LLM не инициализирована, не могу определить категорию.")
        return "Разное" # Категория по умолчанию при ошибке

    # Используем только начало текста для экономии токенов LLM
    text_sample = text[:1500]
    if not text_sample.strip(): # Если текст пустой
        return "Разное"

    prompt = f"""
Определи наиболее подходящую категорию для следующего текста из списка: {CATEGORIES}.
Учитывай также имя файла: {filename}.
Отвечай ТОЛЬКО ОДНИМ СЛОВОМ - названием категории из списка.

Текст (фрагмент):
"{text_sample}"

Категория:
"""
    try:
        category = llm.invoke(prompt).strip().replace(".", "") # Убираем точки на всякий случай
        # Проверяем, есть ли ответ в списке категорий
        if category in CATEGORIES:
            logger.info(f"Категория для '{filename}' определена LLM как: {category}")
            return category
        else:
            logger.warning(f"LLM вернула неизвестную категорию '{category}' для '{filename}'. Используется 'Разное'.")
            return "Разное"
    except Exception as e:
        logger.error(f"Ошибка LLM при определении категории для '{filename}': {e}")
        return "Разное" # Категория по умолчанию при ошибке

# --- Эндпоинт API ---

# @router.post( # Используем APIRouter, если он есть, или app.post
#     "/upload",
#     status_code=status.HTTP_207_MULTI_STATUS, # По умолчанию Multi-Status, т.к. может быть частичный успех
#     summary="Загрузить файлы в чат",
#     response_description="Результаты обработки каждого файла"
# )
async def upload_files_to_chat(
    session_id: str = Form(..., description="ID чата (общего или пользовательского), в который загружаются файлы."),
    files: List[UploadFile] = File(..., description="Один или несколько файлов для загрузки."),
    db: Session = Depends(get_db)
):
    """
    Принимает файлы и ID чата. Обрабатывает каждый файл:
    1. Извлекает текст.
    2. Определяет категорию.
    3. Разбивает на чанки.
    4. Создает эмбеддинги для чанков.
    5. Загружает точки (вектор + payload) в соответствующую коллекцию Qdrant.
    6. Сохраняет метаданные о файле в БД.
    Возвращает статус обработки для каждого файла.
    """
    if not files:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не предоставлены файлы для загрузки."
        )

    # --- Получаем объекты Qdrant и Embeddings ---
    try:
        qdrant = get_qdrant_client()
        embed = get_embedding_model()
    except Exception as init_e:
        logger.critical(f"Критическая ошибка инициализации Qdrant/Embeddings при загрузке: {init_e}", exc_info=True)
        raise HTTPException(status_code=503, detail="Сервис временно недоступен (ошибка базы векторов).")

    # --- Определяем целевую коллекцию Qdrant ---
    try:
        if session_id == GENERAL_CHAT_SESSION_ID:
            target_collection = GENERAL_QDRANT_COLLECTION
            # Дополнительно можно проверить существование коллекции через Qdrant API, если нужно
            # client.get_collection(target_collection)
        else:
            chat = db.query(Chat).filter(Chat.id == session_id).first()
            if not chat:
                # Если чат не найден, не позволяем загружать файлы "в никуда"
                raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail=f"Чат с ID {session_id} не найден. Невозможно загрузить файлы.")
            target_collection = chat.qdrant_collection_name
            if not target_collection:
                 # Этого не должно быть, если чат создан корректно
                 logger.error(f"Чат {session_id} найден, но не имеет имени коллекции Qdrant!")
                 raise HTTPException(status_code=500, detail=f"Ошибка конфигурации чата {session_id}: отсутствует имя коллекции.")

        logger.info(f"Загрузка файлов инициирована в чат ID: {session_id}, целевая коллекция Qdrant: '{target_collection}'")

    except HTTPException as http_exc:
        raise http_exc # Пробрасываем 404 или 500
    except Exception as e:
         logger.error(f"Неожиданная ошибка при определении чата/коллекции для загрузки ({session_id}): {e}", exc_info=True)
         raise HTTPException(status_code=500, detail="Внутренняя ошибка сервера при подготовке к загрузке.")


    results = [] # Список для хранения результатов обработки каждого файла
    docs_to_commit = [] # Список метаданных документов для коммита в БД

    # --- Обработка каждого файла ---
    for file in files:
        file_info = {"filename": file.filename, "status": "error", "error": "Неизвестная ошибка"}
        points_to_upsert: List[PointStruct] = [] # Точки для загрузки в Qdrant для текущего файла
        file_category = "Разное" # Категория по умолчанию
        extracted_text = ""
        file_data = {}

        try:
            # 1. Извлечение текста и базовой информации
            file_data = await process_file(file) # file будет закрыт внутри process_file
            extracted_text = file_data["text"]

            # Если текст не извлечен, файл пропускается, но помечается как обработанный (без ошибки)
            if not extracted_text:
                 logger.warning(f"Текст не извлечен из файла '{file.filename}'. Файл пропущен для индексации.")
                 file_info.update({
                     "status": "skipped",
                     "details": {"message": "Текст не извлечен (возможно, изображение или пустой файл)."},
                     "error": None
                 })
                 results.append(file_info)
                 continue # Переходим к следующему файлу

            # 2. Определение категории
            file_category = detect_category(extracted_text, file.filename)

            # 3. Разбиение на чанки и создание эмбеддингов/точек
            upload_date_iso = datetime.utcnow().isoformat() # Единое время загрузки для всех чанков файла
            chunk_index = 0
            # Используем генератор chunk_text
            for chunk in chunk_text(extracted_text):
                if not chunk.strip(): continue # Пропускаем пустые чанки

                try:
                    # Получаем эмбеддинг для чанка
                    # embedding = await embed.aembed_query(chunk) # Асинхронный вариант
                    embedding = embed.embed_query(chunk) # Синхронный вариант

                    # Создаем точку для Qdrant
                    point = PointStruct(
                        id=str(uuid.uuid4()), # Уникальный ID для каждой точки/чанка
                        vector=embedding,
                        payload={
                            "text": chunk, # Сам текст чанка
                            "source": file_data["original_filename"], # Имя исходного файла
                            "category": file_category, # Определенная категория
                            "upload_date": upload_date_iso, # Дата загрузки (ISO строка)
                            "chunk_index": chunk_index # Порядковый номер чанка в файле
                        }
                    )
                    points_to_upsert.append(point)
                    chunk_index += 1

                except Exception as emb_e:
                    logger.error(f"Ошибка создания эмбеддинга/точки для чанка {chunk_index} файла '{file.filename}': {emb_e}", exc_info=True)
                    # Решаем: пропустить чанк или весь файл? Пропустим чанк.
                    continue

            # Проверяем, были ли созданы точки
            if not points_to_upsert:
                 # Это может случиться, если текст был, но чанки не создались или эмбеддинги не сработали
                 logger.warning(f"Не создано ни одной точки для индексации файла '{file.filename}' (после извлечения текста).")
                 # Не считаем это фатальной ошибкой для всего запроса, но для файла - да.
                 raise FileUploadError("Не удалось создать векторные представления для чанков файла.")

            # 4. Загрузка точек в Qdrant (батчем для одного файла)
            try:
                qdrant.upsert(
                    collection_name=target_collection,
                    points=points_to_upsert,
                    wait=True # Ожидаем подтверждения от Qdrant для надежности
                )
                logger.info(f"Успешно загружено {len(points_to_upsert)} точек для файла '{file.filename}' в коллекцию '{target_collection}'.")
            except Exception as qdrant_e:
                logger.error(f"Ошибка загрузки точек в Qdrant для файла '{file.filename}' (коллекция: {target_collection}): {qdrant_e}", exc_info=True)
                # Проверяем, существует ли коллекция, на всякий случай
                try:
                     qdrant.get_collection(collection_name=target_collection)
                except Exception as get_coll_e:
                     if "not found" in str(get_coll_e).lower():
                         logger.critical(f"Критическая ошибка: Целевая коллекция '{target_collection}' не найдена в Qdrant во время загрузки!")
                         # Возможно, чат был удален во время загрузки?
                         raise FileUploadError(f"Целевая коллекция '{target_collection}' не найдена.")
                     else:
                         # Другая ошибка при проверке коллекции
                         raise FileUploadError(f"Ошибка векторной базы данных при загрузке: {qdrant_e}")
                else:
                     # Коллекция существует, значит проблема в другом (данные, соединение и т.д.)
                     raise FileUploadError(f"Ошибка векторной базы данных при загрузке: {qdrant_e}")


            # 5. Подготовка метаданных файла для сохранения в БД
            doc_metadata = DbDocument(
                file_name=file_data["original_filename"],
                file_type=file_data["file_type"],
                upload_date=datetime.fromisoformat(upload_date_iso), # Преобразуем обратно в datetime
                category=file_category,
                chunks_count=len(points_to_upsert)
                # chat_id=session_id # Раскомментировать, если добавили поле в модель Document
            )
            docs_to_commit.append(doc_metadata) # Добавляем в список для коммита

            # Обновляем информацию о результате для этого файла
            file_info.update({
                "status": "success",
                "details": {
                    "message": "Файл успешно обработан и проиндексирован.",
                    "chunks_indexed": len(points_to_upsert),
                    "detected_category": file_category,
                    "file_size": file_data["file_size"],
                    "file_type": file_data["file_type"],
                },
                "error": None
            })

        # Обработка специфичных ошибок загрузки/обработки файла
        except FileUploadError as e:
            logger.error(f"Ошибка обработки файла '{file.filename}' для чата {session_id}: {e}")
            file_info["error"] = str(e)
        # Обработка других неожиданных ошибок
        except Exception as e:
            logger.error(f"Неожиданная ошибка при обработке файла '{file.filename}' для чата {session_id}: {e}", exc_info=True)
            file_info["error"] = "Внутренняя ошибка сервера при обработке этого файла."
        finally:
            # Добавляем результат обработки файла в общий список
            results.append(file_info)
            # Важно: Не закрываем файл здесь, т.к. он читается и закрывается в process_file

    # --- Сохранение метаданных всех успешно обработанных файлов в БД ---
    if docs_to_commit:
        try:
            db.add_all(docs_to_commit)
            db.commit()
            logger.info(f"Успешно сохранено {len(docs_to_commit)} записей метаданных файлов в БД.")
        except Exception as db_e:
             logger.error(f"Критическая ошибка сохранения метаданных файлов в БД: {db_e}", exc_info=True)
             db.rollback()
             # Как обработать эту ошибку? Файлы уже в Qdrant.
             # Можно попытаться удалить загруженные точки из Qdrant или просто сообщить об ошибке.
             # Пометим все успешные загрузки как имеющие проблемы с метаданными.
             for res in results:
                 if res["status"] == "success":
                     res["status"] = "warning"
                     res["details"]["message"] = "Файл проиндексирован, но произошла ошибка сохранения метаданных."
                     res["error"] = "Ошибка БД при сохранении информации о файле."

    # --- Определение финального статуса ответа ---
    successful_uploads = [r for r in results if r["status"] == "success"]
    failed_uploads = [r for r in results if r["status"] == "error"]
    skipped_uploads = [r for r in results if r["status"] == "skipped"]
    warning_uploads = [r for r in results if r["status"] == "warning"]

    final_status_code = status.HTTP_207_MULTI_STATUS # По умолчанию
    response_message = "Обработка файлов завершена."

    if successful_uploads and not failed_uploads and not skipped_uploads and not warning_uploads:
        final_status_code = status.HTTP_201_CREATED # Все успешно
        response_message = "Все файлы успешно загружены и проиндексированы."
    elif not successful_uploads and failed_uploads:
        final_status_code = status.HTTP_422_UNPROCESSABLE_ENTITY # Все с ошибками
        response_message = "Не удалось обработать ни один из файлов."
    # Другие комбинации остаются HTTP_207_MULTI_STATUS

    return JSONResponse(
        status_code=final_status_code,
        content={
            "message": response_message,
            "results": results # Полный список результатов для каждого файла
        }
    )

def get_categories():
    """Возвращает список доступных категорий."""
    # Можно добавить логику получения категорий из БД или другого источника
    return JSONResponse(content={"categories": CATEGORIES})

# --- Роутер (если используется) ---
# Пример использования APIRouter, если ваш проект структурирован так
from fastapi import APIRouter
router = APIRouter(prefix="/api", tags=["Files & Categories"]) # Пример

# Добавляем эндпоинты в роутер
# router.post("/upload")(upload_files_to_chat) # Заменено выше
router.get("/categories")(get_categories)