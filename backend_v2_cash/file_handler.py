from fastapi import UploadFile, HTTPException, Depends, File, status
from fastapi.responses import JSONResponse
from sqlalchemy.orm import Session
import os
import fitz
from docx import Document
import uuid
import logging
from datetime import datetime
from typing import List, Dict, Optional
from langchain_ollama import OllamaLLM
from pathlib import Path
import chardet

from qdrant_db import qdrant_client, collection_name, embed_model
from database import get_db, Document as DbDocument
from chat_AI import CATEGORIES

logger = logging.getLogger(__name__)
UPLOAD_DIR = "../backend_v2/uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

SUPPORTED_EXTENSIONS = {
    '.pdf': 'application/pdf',
    '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    '.txt': 'text/plain'
}

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
llm = OllamaLLM(model="llama3.2")


class FileUploadError(Exception):
    """Кастомное исключение для ошибок загрузки файлов"""
    pass


def get_file_encoding(file_path: str) -> str:
    """Определяет кодировку текстового файла"""
    with open(file_path, 'rb') as f:
        result = chardet.detect(f.read(10000))  # Проверяем первые 10KB файла
    return result['encoding']


def chunk_text(text: str, chunk_size: int = 512, overlap: int = 50) -> List[str]:
    """Разбивает текст на перекрывающиеся чанки"""
    sentences = [s.strip() for s in text.split('. ') if s.strip()]
    chunks = []
    current_chunk = []
    current_length = 0

    for sentence in sentences:
        sentence_length = len(sentence.split())
        if current_length + sentence_length > chunk_size and current_chunk:
            chunks.append('. '.join(current_chunk) + '.')
            current_chunk = []
            current_length = 0

        current_chunk.append(sentence)
        current_length += sentence_length

    if current_chunk:
        chunks.append('. '.join(current_chunk) + '.')

    # Добавляем перекрытие между чанками
    if len(chunks) > 1 and overlap > 0:
        for i in range(1, len(chunks)):
            prev_chunk = chunks[i - 1].split('. ')
            overlap_sentences = prev_chunk[-overlap:] if len(prev_chunk) > overlap else prev_chunk
            chunks[i] = '. '.join(overlap_sentences) + '. ' + chunks[i]

    return chunks


async def process_file(file: UploadFile) -> Dict:
    """Обрабатывает загруженный файл и извлекает текст"""
    try:
        # Проверка размера файла
        file.file.seek(0, 2)  # Перемещаемся в конец файла
        file_size = file.file.tell()
        file.file.seek(0)  # Возвращаемся в начало

        if file_size > MAX_FILE_SIZE:
            raise FileUploadError(f"Файл слишком большой. Максимальный размер: {MAX_FILE_SIZE // (1024 * 1024)}MB")

        file_ext = Path(file.filename).suffix.lower()
        logger.info(f"Обработка файла: {file.filename}, размер: {file_size} байт, расширение: {file_ext}")

        if file_ext not in SUPPORTED_EXTENSIONS:
            raise FileUploadError(
                f"Неподдерживаемый формат файла: {file_ext}. "
                f"Поддерживаются: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
            )

        # Сохраняем файл временно
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = os.path.join(UPLOAD_DIR, unique_filename)

        with open(file_path, "wb") as buffer:
            content = await file.read()
            if not content:
                raise FileUploadError("Файл пустой")
            buffer.write(content)

        # Обработка в зависимости от типа файла
        text = ""
        try:
            if file_ext == '.pdf':
                with fitz.open(file_path) as doc:
                    text = "\n".join(page.get_text() for page in doc)
            elif file_ext == '.docx':
                doc = Document(file_path)
                text = "\n".join(para.text for para in doc.paragraphs)
            elif file_ext == '.txt':
                encoding = get_file_encoding(file_path)
                with open(file_path, "r", encoding=encoding) as f:
                    text = f.read()
        except Exception as e:
            raise FileUploadError(f"Ошибка чтения файла: {str(e)}")
        finally:
            if os.path.exists(file_path):
                os.remove(file_path)

        if not text.strip():
            raise FileUploadError("Не удалось извлечь текст из файла (пустое содержимое)")

        logger.info(f"Успешно извлечен текст, длина: {len(text)} символов")
        return {
            "text": text,
            "original_filename": file.filename,
            "saved_filename": unique_filename,
            "file_type": SUPPORTED_EXTENSIONS[file_ext],
            "file_size": file_size
        }

    except Exception as e:
        logger.error(f"Ошибка обработки файла: {str(e)}", exc_info=True)
        raise FileUploadError(str(e))


def detect_category(text: str) -> str:
    """Определяет категорию текста с помощью LLM"""
    prompt = f"""
    Определи категорию текста из списка: {CATEGORIES}
    Отвечай ТОЛЬКО названием категории без пояснений.
    Текст: "{text[:1000]}"
    Категория:
    """
    try:
        category = llm.invoke(prompt).strip()
        return category if category in CATEGORIES else "Разное"
    except Exception as e:
        logger.error(f"Ошибка определения категории: {e}")
        return "Разное"


async def upload_file(
    files: List[UploadFile] = File(..., description="Файлы для загрузки"),
    db: Session = Depends(get_db)
):
    """Эндпоинт для загрузки одного или нескольких файлов"""
    if not files:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Не предоставлены файлы для загрузки"
        )

    results = []
    for file in files:
        try:
            file_data = await process_file(file)
            text = file_data["text"]
            upload_date = datetime.utcnow()

            category = detect_category(text)
            logger.info(f"Файл '{file.filename}' обработан, категория: {category}")

            chunks = chunk_text(text)
            if not chunks:
                raise FileUploadError("Не удалось разбить текст на фрагменты")

            # Создаем векторные представления
            points = []
            for i, chunk in enumerate(chunks):
                try:
                    embedding = embed_model.embed_query(chunk)
                    points.append({
                        "id": str(uuid.uuid4()),
                        "vector": embedding,
                        "payload": {
                            "text": chunk,
                            "source": file_data["original_filename"],
                            "category": category,
                            "upload_date": upload_date.isoformat(),
                            "chunk_index": i
                        }
                    })
                except Exception as e:
                    logger.error(f"Ошибка создания эмбеддинга: {str(e)}")
                    continue

            if not points:
                raise FileUploadError("Не удалось создать векторные представления")

            # Сохраняем в Qdrant
            try:
                qdrant_client.upsert(
                    collection_name=collection_name,
                    points=points,
                    wait=True
                )
            except Exception as e:
                raise FileUploadError(f"Ошибка векторной БД: {str(e)}")

            # Сохраняем метаданные
            doc = DbDocument(
                file_name=file_data["original_filename"],
                stored_name=file_data["saved_filename"],
                file_type=file_data["file_type"],
                upload_date=upload_date,
                category=category,
                vector_id=str(uuid.uuid4()),
                chunks_count=len(points))
            db.add(doc)
            db.commit()

            results.append({
                "status": "success",
                "filename": file_data["original_filename"],
                "details": {
                    "chunks": len(points),
                    "category": category,
                    "file_size": file_data["file_size"],
                    "file_type": file_data["file_type"]
                }
            })

        except FileUploadError as e:
            logger.error(f"Ошибка загрузки файла {file.filename}: {str(e)}")
            results.append({
                "status": "error",
                "filename": file.filename,
                "error": str(e)
            })
        except Exception as e:
            logger.error(f"Неожиданная ошибка при обработке {file.filename}: {str(e)}")
            results.append({
                "status": "error",
                "filename": file.filename,
                "error": "Внутренняя ошибка сервера"
            })

    # Проверяем, есть ли успешные загрузки
    if any(r["status"] == "success" for r in results):
        return JSONResponse(
            status_code=status.HTTP_201_CREATED,
            content={"results": results}
        )
    else:
        raise HTTPException(
            status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
            detail={"results": results}
        )

def get_categories():
    """Возвращает список доступных категорий"""
    return JSONResponse(content={"categories": CATEGORIES})